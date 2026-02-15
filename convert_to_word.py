"""
Convert final_mma_extracted.md to Word document
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def parse_markdown_to_docx(md_file, docx_file):
    """Convert markdown file to Word document with formatting"""
    #piyush
    # Read the markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create a new Document
    doc = Document()
    
    # Process line by line
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        # Skip empty lines at the start
        if not line.strip():
            i += 1
            continue
        
        # Headers
        if line.startswith('# '):
            p = doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith('## '):
            p = doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('### '):
            p = doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('#### '):
            p = doc.add_heading(line[5:].strip(), level=4)
        
        # Code blocks
        elif line.strip().startswith('```'):
            # Collect code block
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            
            # Add code block with monospace font
            if code_lines:
                p = doc.add_paragraph()
                run = p.add_run('\n'.join(code_lines))
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                p.paragraph_format.left_indent = Inches(0.5)
        
        # Bold text patterns
        elif '**' in line:
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
        
        # Bullet points
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            # Handle bold within bullets
            p = doc.add_paragraph(style='List Bullet')
            if '**' in text:
                parts = re.split(r'(\*\*.*?\*\*)', text)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    else:
                        run = p.add_run(part)
            else:
                p.add_run(text)
        
        # Numbered lists
        elif re.match(r'^\d+\.', line.strip()):
            text = re.sub(r'^\d+\.\s*', '', line.strip())
            p = doc.add_paragraph(text, style='List Number')
        
        # Regular paragraph
        elif line.strip():
            # Handle inline formatting
            if '**' in line:
                p = doc.add_paragraph()
                parts = re.split(r'(\*\*.*?\*\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    else:
                        p.add_run(part)
            else:
                doc.add_paragraph(line)
        
        i += 1
    
    # Save the document
    doc.save(docx_file)
    print(f"✓ Converted {md_file} to {docx_file}")

if __name__ == "__main__":
    input_file = "final_mma_extracted.md"
    output_file = "final_mma_extracted.docx"
    
    try:
        parse_markdown_to_docx(input_file, output_file)
        print(f"\n✓ Successfully created Word document: {output_file}")
    except Exception as e:
        print(f"✗ Error: {e}")
